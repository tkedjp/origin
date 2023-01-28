import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd

# wordファイルを新規作成
doc = docx.Document()

#ファイルの取得
df = pd.read_csv('list.csv')  

for i, r in df.iterrows():
    # print(r['ホテル名'], r['詳細ページ'])
    if i == '':
        break

    #表題
    doc.add_heading(r['ホテル名'], 0)
    doc.add_paragraph('住所：'+ r['住所'])

    #部屋タイプ内訳をテーブルにする
    table = doc.add_table(rows=1, cols=5)

    header_cells = table.rows[0].cells
    header_cells[0].text = 'シングル'
    header_cells[1].text = 'ダブル'
    header_cells[2].text = 'ツイン'
    header_cells[3].text = 'スイート'
    header_cells[4].text = '総部屋数'

    row_cells = table.add_row().cells
    row_cells[0].text = (str(r['シングル']))
    row_cells[1].text = (str(r['ダブル']))
    row_cells[2].text = (str(r['ツイン']))
    row_cells[3].text = (str(r['スイート']))
    row_cells[4].text = (str(r['総部屋数']))

    doc.add_paragraph('')
    doc.add_paragraph('室料：'+ str(r['室料']))
    doc.add_paragraph('1人あたり：'+ str(r['1人あたり']))
    doc.add_paragraph('駐車場：'+ r['駐車場'].strip())
    doc.add_paragraph('詳細ページ：'+ r['詳細ページ'])

    # 改ページ
    doc.add_page_break()

# すべての行を左揃えにする
for p in doc.paragraphs:
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
# ファイルを保存
doc.save('hotel_list.docx')