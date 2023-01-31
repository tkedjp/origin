import tkinter
from tkinter import ttk
import datetime
from datetime import timedelta

from time import sleep

import requests
from bs4 import BeautifulSoup
import pandas as pd

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

# ウィンドウの作成
root = tkinter.Tk()
root.title('じゃらん')
root.geometry('380x210')
root.resizable(0, 0)

def save():
    date_box_value = date_box.get()
    month = date_box_value[5:7]
    day = date_box_value[8:10]
    stay_count = stay_count_box.get()
    room_count = room_count_box.get()
    adult_num = adult_num_box.get()

    hotel_list = []
    base_url = 'https://www.jalan.net/040000/LRG_040200/SML_040202/?screenId=UWW1402&distCd=01&listId=0&activeSort=0&mvTabFlg=1&stayYear=2023&stayMonth=' + month + '&stayDay=' + day + '&stayCount=' + stay_count + '&roomCount=' + room_count +'&adultNum=' + adult_num +'&yadHb=1&roomCrack=200000&kenCd=040000&lrgCd=040200&smlCd=040202&vosFlg=6&idx={}'

    sleep(3)

    r = requests.get(base_url, timeout=7.5)
    if r.status_code >= 400:
        print(F'{base_url}は無効です')

    soup = BeautifulSoup(r.content, 'lxml')

    number = soup.select_one('td.jlnpc-planListCnt-header > span.s16_F60b').text
    max_page_index = int(number) // 30 + 1

    for i in range(max_page_index):
        url = base_url.format(30*i)

        sleep(3)

        page_r = requests.get(url, timeout=7.5)
        if page_r.status_code >= 400:
            print(F'{url}は無効です')
            continue

        page_soup = BeautifulSoup(page_r.content, 'lxml')

        #最安料金と1人あたりの料金
        table_soup = page_soup.select('div.p-yadoCassette__body.p-searchResultItem__body')
        for table in table_soup:
            hotel = table.select_one('a > div > div > div.p-searchResultItem__summaryInner > div.p-searchResultItem__summaryLeft > h2').text
            room_price = table.select_one('a > div > div > div.p-searchResultItem__summaryInner > div.p-searchResultItem__summaryRight > dl > dd > span.p-searchResultItem__lowestPriceValue').text
            per_price = table.select_one('a > div > div > div.p-searchResultItem__summaryInner > div.p-searchResultItem__summaryRight > dl > dd > span.p-searchResultItem__lowestUnitPrice').text
            page_urls = table.select('a.jlnpc-yadoCassette__link')

            for i, page_url in enumerate(page_urls):
                page_url = 'https://www.jalan.net' + page_url.get('href')

                sleep(3)

                hotel_page_r = requests.get(page_url, timeout=7.5)
                if hotel_page_r.status_code >= 400:
                    print(F'{page_url}は無効です')
                    continue
            
                #ホテル名
                hotel_page_soup = BeautifulSoup(hotel_page_r.content, 'lxml')

                #住所
                address_tags = hotel_page_soup.select_one('#jlnpc-main-contets-area > div.shisetsu-accesspartking_body_wrap > table tr:nth-child(1) > td').text
                if address_tags == None:
                    continue

                else: 
                    address = address_tags.replace('大きな地図をみる', '')
                    address = address.strip()
            
                #駐車場
                parking_tags = hotel_page_soup.select_one('#jlnpc-main-contets-area > div.shisetsu-accesspartking_body_wrap > table tr:nth-child(3) > td').text  
                if parking_tags == None:
                    continue

                else:
                    parking = parking_tags.replace('\n','')
                    parking = parking.strip()

                #タイプ別の室数
                room_tag = hotel_page_soup.select_one('.shisetsu-roomsetsubi_body')
                tags = room_tag.text

                if '総部屋数' not in tags:
                    single = room_tag.select_one('tr:nth-child(2) > td > div > table tr:nth-child(2) > td:first-child').text
                    double = room_tag.select_one('tr:nth-child(2) > td > div > table tr:nth-child(2) > td:nth-child(2)').text
                    twin = room_tag.select_one('tr:nth-child(2) > td > div > table tr:nth-child(2) > td:nth-child(3)').text
                    sweet = room_tag.select_one('tr:nth-child(2) > td > div > table tr:nth-child(2) > td:last-child').text
                    total = None

                elif 'シングル' not in tags:
                    single = None
                    double = None
                    twin = None
                    sweet = None
                    total = room_tag.select_one('tr:nth-child(1) > td > div > table tr:nth-child(2) > td:nth-child(5)').text
                    total = total.strip()

                else:
                    single = room_tag.select_one('tr:nth-child(3) > td > div > table tr:nth-child(2) > td:first-child').text
                    double = room_tag.select_one('tr:nth-child(3) > td > div > table tr:nth-child(2) > td:nth-child(2)').text
                    twin = room_tag.select_one('tr:nth-child(3) > td > div > table tr:nth-child(2) > td:nth-child(3)').text
                    sweet = room_tag.select_one('tr:nth-child(3) > td > div > table tr:nth-child(2) > td:last-child').text
                    total = room_tag.select_one('tr:nth-child(1) > td > div > table tr:nth-child(2) > td:nth-child(5)').text
                    total = total.strip()

                hotel_list.append({
                    'ホテル名': hotel,
                    '詳細ページ': page_url,
                    '住所': address,
                    'シングル': single,
                    'ダブル': double,
                    'ツイン': twin,
                    'スイート': sweet,
                    '総部屋数': total,
                    '室料': room_price,
                    '1人あたり': per_price,
                    '駐車場': parking
                })
                print(hotel_list[-1])

    #csv出力
    df = pd.DataFrame(hotel_list)
    df.to_csv('list.csv', index=False, encoding='utf-8-sig')

    # wordファイルを新規作成
    doc = docx.Document()

    #ファイルの取得
    df = pd.read_csv('list.csv')  

    for i, r in df.iterrows():
        if i == '':
            break

        #表題
        doc.add_heading(r['ホテル名'], 0)
        doc.add_paragraph('住所：'+ r['住所'])

        #部屋タイプ内訳をテーブルにする
        table = doc.add_table(rows=1, cols=5)

        header_cells = table.rows[0].cells
        header_cells[0].text = 'シングル'
        header_cells[1].text = 'ダブル'
        header_cells[2].text = 'ツイン'
        header_cells[3].text = 'スイート'
        header_cells[4].text = '総部屋数'

        row_cells = table.add_row().cells
        row_cells[0].text = (str(r['シングル']))
        row_cells[1].text = (str(r['ダブル']))
        row_cells[2].text = (str(r['ツイン']))
        row_cells[3].text = (str(r['スイート']))
        row_cells[4].text = (str(r['総部屋数']))

        doc.add_paragraph('')
        doc.add_paragraph('室料：'+ str(r['室料']))
        doc.add_paragraph('1人あたり：'+ str(r['1人あたり']))
        doc.add_paragraph('駐車場：'+ r['駐車場'].strip())
        doc.add_paragraph('詳細ページ：'+ r['詳細ページ'])

        # 改ページ
        doc.add_page_break()

    # すべての行を左揃えにする
    for p in doc.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    # ファイルを保存
    doc.save('hotel_list.docx')

#日付
date = tkinter.Label(text='チェックインする日：').grid(row=1, column=1, padx=5, pady=5)
date_box = ttk.Combobox(values=[datetime.date.today() + datetime.timedelta(days=i) for i in range(180)])
date_box.grid(row=1, column=2, padx=5, pady=5)

#泊数
stay_count = tkinter.Label(text='泊数：').grid(row=2, column=1, padx=5, pady=5)
stay_count_box = ttk.Combobox(values=[i+1 for i in range(10)])
stay_count_box.grid(row=2, column=2, padx=5, pady=5)

#室数
room_count = tkinter.Label(text='室数：').grid(row=3, column=1, padx=5, pady=5)
room_count_box = ttk.Combobox(values=[i+1 for i in range(10)])
room_count_box.grid(row=3, column=2, padx=5, pady=5)

#人数
adult_num = tkinter.Label(text='人数：').grid(row=4, column=1, padx=5, pady=5)
adult_num_box = ttk.Combobox(values=[i+1 for i in range(9)])
adult_num_box.grid(row=4, column=2, padx=5, pady=5)

#実行
save_button = tkinter.Button(text='取得', command=save).grid(row=5, column=2, padx=5, pady=20, ipadx=10, ipady=10)

# ウィンドウのループ処理
root.mainloop()